"""
word_creator — AresAI SkillRegistry

Skill for creating professional Word .docx documents.
Uses the verified docx_helper module — no raw python-docx calls.

Input config (dict):
    title:              str  — document main title
    output_filename:    str  — e.g. 'Report.docx', saved under workspace/
    context_data_path:  str  — default: 'workspace/context_data.json'
    images:             list — [{"path": "workspace/x.jpg", "caption": "..."}]
    theme:              str  — "premium_automotive" | "default"
    sections_override:  list — [{"heading", "type", "data_key"}] (optional)

Returns:
    str — relative path to output file, e.g. 'workspace/Report.docx'
         or error string starting with 'Error:'
"""

import json
import os
import sys
from pathlib import Path

# ── Project root on sys.path so docx_helper is importable ─────────────────────
_PROJECT_ROOT = Path(__file__).parent.parent.parent.parent
sys.path.insert(0, str(_PROJECT_ROOT))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from tools.docx_helper import (
    AUDI_DARK, AUDI_SILVER,
    add_body_text, add_bullet, add_divider,
    add_image_safe, add_premium_heading, add_premium_table,
    set_default_style, set_doc_margins,
)


# ── Internal helpers ───────────────────────────────────────────────────────────

def _flatten_to_rows(data, max_depth: int = 2) -> list[list[str]]:
    """Recursively flatten a dict/list into [Key, Value] rows."""
    rows = []
    if isinstance(data, dict):
        for k, v in data.items():
            label = str(k).replace("_", " ").title()
            if isinstance(v, (dict, list)) and max_depth > 1:
                rows.append([f"▶ {label}", ""])
                rows.extend(_flatten_to_rows(v, max_depth - 1))
            else:
                rows.append([label, _scalar(v)])
    elif isinstance(data, list):
        for i, item in enumerate(data):
            if isinstance(item, dict):
                rows.extend(_flatten_to_rows(item, max_depth))
            else:
                rows.append([str(i + 1), _scalar(item)])
    # NEW: handle plain string
    elif isinstance(data, str):
        rows.append(["", data])
    return rows


def _scalar(v) -> str:
    if isinstance(v, list):
        return " · ".join(str(x) for x in v[:6])
    return str(v)


def _resolve_data_path(context: dict, data_key: str):
    """Resolve simple paths like 'sections[0].content' from context_data.json."""
    if not data_key:
        return None
    if data_key in context:
        return context.get(data_key)

    current = context
    parts = str(data_key).split(".")
    for part in parts:
        if current is None:
            return None
        while "[" in part and part.endswith("]"):
            prefix, _, rest = part.partition("[")
            if prefix:
                if not isinstance(current, dict):
                    return None
                current = current.get(prefix)
            index_text = rest[:-1]
            if not index_text.isdigit() or not isinstance(current, list):
                return None
            index = int(index_text)
            if index >= len(current):
                return None
            current = current[index]
            part = ""
        if part:
            if not isinstance(current, dict):
                return None
            current = current.get(part)
    return current


def _section_specs(doc: Document, data) -> None:
    rows = _flatten_to_rows(data)
    if rows:
        add_premium_table(doc, ["Параметр", "Значення"], rows, col_widths=[2.5, 3.5])
    else:
        add_body_text(doc, "Дані відсутні.")


def _section_bullets(doc: Document, data) -> None:
    items = data if isinstance(data, list) else list(data.values()) if isinstance(data, dict) else [str(data)]
    for item in items:
        add_bullet(doc, str(item))


def _section_comparison(doc: Document, competitors: list) -> None:
    """Build a multi-column comparison table from a list of competitor dicts."""
    if not competitors or not isinstance(competitors, list):
        add_body_text(doc, "Дані для порівняння відсутні.")
        return

    # Gather all unique keys
    all_keys: list[str] = []
    for comp in competitors:
        if isinstance(comp, dict):
            for k in comp:
                if k not in all_keys:
                    all_keys.append(k)

    name_key  = all_keys[0] if all_keys else "model"
    spec_keys = all_keys[1:]

    headers = ["Параметр"] + [
        str(c.get(name_key, f"Model {i+1}")) for i, c in enumerate(competitors)
    ]
    rows: list[list[str]] = []
    for key in spec_keys:
        row = [key.replace("_", " ").title()] + [
            _scalar(c.get(key, "—")) for c in competitors
        ]
        rows.append(row)

    if rows:
        n = len(competitors)
        col_w = [2.0] + [min(4.0 / n, 2.2)] * n
        add_premium_table(doc, headers, rows, col_widths=col_w)


def _section_table(doc: Document, data) -> None:
    if not isinstance(data, dict):
        _section_specs(doc, data)
        return
    headers = data.get("headers") or data.get("columns") or []
    rows = data.get("rows") or []
    if not headers or not rows:
        _section_specs(doc, data)
        return
    add_premium_table(doc, [str(h) for h in headers], [[_scalar(cell) for cell in row] for row in rows])


def _section_offers(doc: Document, data) -> None:
    offers = data if isinstance(data, list) else data.get("offers", []) if isinstance(data, dict) else []
    if not offers:
        add_body_text(doc, "Пропозиції не знайдено.")
        return
    headers = ["Товар", "Продавець", "Ціна", "Посилання", "Чому вибрано"]
    rows = []
    for offer in offers:
        if not isinstance(offer, dict):
            continue
        price = str(offer.get("price", ""))
        currency = str(offer.get("currency", ""))
        rows.append([
            offer.get("product", ""),
            offer.get("seller", ""),
            f"{price} {currency}".strip(),
            offer.get("url", "") or offer.get("link", ""),
            offer.get("why_selected", "") or offer.get("freshness_note", ""),
        ])
    if rows:
        add_premium_table(doc, headers, rows)
    else:
        _section_specs(doc, data)


def _section_pros_cons(doc: Document, data: dict) -> None:
    pros = data.get("pros", [])
    cons = data.get("cons", [])
    add_premium_heading(doc, "✅ Переваги", level=3, font_size=12, color=RGBColor(0x1A, 0x7A, 0x1A))
    for p in pros:
        add_bullet(doc, str(p))
    add_premium_heading(doc, "⚠️ Недоліки", level=3, font_size=12, color=RGBColor(0x9A, 0x3A, 0x00))
    for c in cons:
        add_bullet(doc, str(c))


# ── SECTION_ORDER: auto-detection priority ─────────────────────────────────────
_SECTION_ORDER = [
    ("comparison_table",        "Таблиця порівняння",        "table"),
    ("top_offers",              "Топ пропозицій з посиланнями", "offers"),
    ("product_offers",          "Пропозиції товарів",        "offers"),
    ("technical_specifications", "Технічні характеристики",  "specs"),
    ("interior_features",        "Оснащення інтер'єру",      "bullets"),
    ("exterior_features",        "Зовнішнє оформлення",      "bullets"),
    ("safety_systems",           "Системи безпеки",          "bullets"),
    ("competitors",              "Порівняння з конкурентами", "comparison"),
    ("ukraine_prices",           "Ціни в Україні",           "specs"),
    ("pricing",                  "Ціни та комплектації",     "specs"),
    ("pros_and_cons",            "Плюси та мінуси",          "pros_cons"),
    ("warranty",                 "Гарантія та обслуговування","specs"),
    ("performance",              "Динамічні характеристики", "specs"),
    ("fuel_efficiency",          "Паливна економічність",    "specs"),
]


# ── Main execute() ─────────────────────────────────────────────────────────────

def execute(config: dict, workdir: str = ".") -> str:
    """Create a professional Word document from the given config and context data."""
    os.makedirs(os.path.join(workdir, "workspace"), exist_ok=True)

    title           = config.get("title", "Document")
    output_filename = config.get("output_filename", "output.docx")
    ctx_rel         = config.get("context_data_path", "workspace/context_data.json")
    images_conf     = config.get("images", [])
    images_disabled = bool(config.get("images_disabled", False))
    sections_over   = config.get("sections_override")

    # ── Load context data ──────────────────────────────────────────────────────
    ctx_abs = os.path.join(workdir, ctx_rel)
    context: dict = {}
    if os.path.exists(ctx_abs):
        try:
            with open(ctx_abs, "r", encoding="utf-8") as f:
                context = json.load(f)
        except Exception as e:
            print(f"[word_creator] Warning: could not read context_data: {e}")

    # ── Create document ────────────────────────────────────────────────────────
    doc = Document()
    set_doc_margins(doc)
    set_default_style(doc)

    # Cover heading
    add_premium_heading(doc, title, level=1, font_size=26,
                        align=WD_ALIGN_PARAGRAPH.CENTER, space_before=36, space_after=8)
    add_divider(doc)

    # Hero image (first in list)
    if images_conf and not images_disabled:
        hero     = images_conf[0]
        hero_abs = os.path.join(workdir, hero.get("path", ""))
        add_image_safe(doc, hero_abs, width=5.5, caption=hero.get("caption", ""))
        doc.add_paragraph()

    # Executive summary / description (shown after title)
    description = context.get("description", "")
    if description:
        add_body_text(doc, description)
        doc.add_paragraph()

    # ── Sections ───────────────────────────────────────────────────────────────
    def _render_section(heading: str, render_type: str, data) -> None:
        if not data:
            return
        add_premium_heading(doc, heading, level=2, font_size=14, space_before=14, space_after=4)
        add_divider(doc)
        
        if isinstance(data, str) and data.strip():
            add_body_text(doc, data)
            doc.add_paragraph()
            return

        if render_type == "specs":
            _section_specs(doc, data)
        elif render_type == "bullets":
            _section_bullets(doc, data)
        elif render_type == "comparison":
            _section_comparison(doc, data)
        elif render_type == "table":
            _section_table(doc, data)
        elif render_type == "offers":
            _section_offers(doc, data)
        elif render_type == "pros_cons":
            _section_pros_cons(doc, data)
        else:
            _section_specs(doc, data)
        doc.add_paragraph()

    if sections_over:
        # Explicit sections from Blacksmith config
        for sec in sections_over:
            h    = sec.get("heading", "Section")
            key  = sec.get("data_key", "")
            stype = sec.get("type", "specs")
            _render_section(h, stype, _resolve_data_path(context, key))
    else:
        rendered = set()
        rendered.add("title")       # already used as doc heading
        rendered.add("description") # already rendered above
        rendered.add("evidence_profile")
        rendered.add("data_quality")
        
        # ── Option C: Universal sections list ────────────────────────────────
        sections_list = context.get("sections", [])
        if sections_list and isinstance(sections_list, list):
            rendered.add("sections")
            for section in sections_list:
                if not isinstance(section, dict):
                    continue
                heading = section.get("title", "")
                content = section.get("content", "")
                if heading and content:
                    add_premium_heading(doc, heading, level=2, font_size=14,
                                        space_before=14, space_after=4)
                    add_divider(doc)
                    add_body_text(doc, content)
                    doc.add_paragraph()
                    
        # ── Domain-specific keys from priority order ─────────────────────────
        for data_key, heading, render_type in _SECTION_ORDER:
            if data_key in context and data_key not in rendered:
                rendered.add(data_key)
                _render_section(heading, render_type, context[data_key])
                
        # ── Render any remaining unknown keys ────────────────────────────────
        for key, val in context.items():
            if key in rendered:
                continue
            if "iteration_" in key or not val:
                continue
            heading = key.replace("_", " ").title()
            add_premium_heading(doc, heading, level=2, font_size=14,
                                space_before=14, space_after=4)
            add_divider(doc)
            _section_specs(doc, val)
            doc.add_paragraph()

    # ── Additional gallery images ──────────────────────────────────────────────
    remaining_imgs = images_conf[1:]
    if remaining_imgs:
        add_premium_heading(doc, "Фотогалерея", level=2, font_size=14, space_before=14, space_after=4)
        add_divider(doc)
        for img_info in remaining_imgs:
            img_abs = os.path.join(workdir, img_info.get("path", ""))
            add_image_safe(doc, img_abs, width=4.5, caption=img_info.get("caption", ""))

    # ── Save ───────────────────────────────────────────────────────────────────
    output_filename = str(output_filename).replace("\\", "/").lstrip("/")
    if output_filename.startswith("workspace/"):
        output_filename = output_filename.split("/", 1)[1]
    output_rel = os.path.join("workspace", output_filename)
    output_abs = os.path.join(workdir, output_rel)
    os.makedirs(os.path.dirname(output_abs), exist_ok=True)
    doc.save(output_abs)
    print(f"[word_creator] ✅ Document saved: {output_rel}")
    return output_rel
