"""
docx_helper.py — Verified python-docx helpers for AresAI generated agents.

Import this in any generated Word-document script:
    from tools.docx_helper import add_premium_heading, add_premium_table, add_image_safe, set_doc_theme

All functions use correct python-docx API (tested). Never write raw python-docx calls in
generated code — always prefer these helpers to avoid AttributeError / style mismatches.
"""

import os
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor


# ── Theme colours ─────────────────────────────────────────────────────────────

AUDI_DARK   = RGBColor(0x1A, 0x1A, 0x1A)   # Near-black
AUDI_SILVER = RGBColor(0xBB, 0xBB, 0xBB)   # Silver accent
AUDI_RED    = RGBColor(0xBB, 0x00, 0x00)   # Audi red
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)


# ── Document setup ─────────────────────────────────────────────────────────────

def set_doc_margins(doc: Document, top=0.7, bottom=0.7, left=0.9, right=0.9) -> None:
    """Set page margins in inches."""
    for section in doc.sections:
        section.top_margin    = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin   = Inches(left)
        section.right_margin  = Inches(right)


def set_default_style(doc: Document, font_name: str = "Calibri", font_size: int = 11) -> None:
    """Set the Normal style font for the whole document."""
    style = doc.styles["Normal"]
    style.font.name     = font_name
    style.font.size     = Pt(font_size)
    style.font.color.rgb = AUDI_DARK


# ── Headings ──────────────────────────────────────────────────────────────────

def add_premium_heading(
    doc: Document,
    text: str,
    level: int = 1,
    color: RGBColor = AUDI_DARK,
    font_size: int = 20,
    bold: bool = True,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    space_before: float = 12,
    space_after: float = 6,
) -> None:
    """Add a styled heading paragraph."""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold       = bold
    run.font.size  = Pt(font_size)
    run.font.color.rgb = color
    run.font.name  = "Calibri"


def add_divider(doc: Document, color: RGBColor = AUDI_SILVER) -> None:
    """Add a thin horizontal rule paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    # Use paragraph border as divider
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    r, g, b = int(color[0]), int(color[1]), int(color[2])
    bottom.set(qn("w:color"), f"{r:02X}{g:02X}{b:02X}")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── Tables ────────────────────────────────────────────────────────────────────

def add_premium_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    header_bg: RGBColor = AUDI_DARK,
    header_fg: RGBColor = WHITE,
    alt_row_bg: RGBColor = RGBColor(0xF5, 0xF5, 0xF5),
    col_widths: list[float] | None = None,
) -> None:
    """
    Add a styled table with header row and alternating row shading.

    Args:
        headers:   Column header strings
        rows:      List of row data (each row is a list of strings)
        col_widths: Column widths in inches. If None, equally distributed across 6".
    """
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Table Grid"

    # Default equal column widths
    if col_widths is None:
        col_widths = [6.0 / n_cols] * n_cols

    # Set column widths
    for i, width in enumerate(col_widths):
        for cell in table.columns[i].cells:
            cell.width = Inches(width)

    def _set_cell_bg(cell, color: RGBColor) -> None:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  f"{int(color[0]):02X}{int(color[1]):02X}{int(color[2]):02X}")
        tcPr.append(shd)

    # Header row
    header_row = table.rows[0]
    for idx, (cell, hdr) in enumerate(zip(header_row.cells, headers)):
        _set_cell_bg(cell, header_bg)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run  = para.add_run(hdr)
        run.bold           = True
        run.font.color.rgb = header_fg
        run.font.name      = "Calibri"
        run.font.size      = Pt(10)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg  = alt_row_bg if r_idx % 2 == 0 else WHITE
        for c_idx, (cell, value) in enumerate(zip(row.cells, row_data)):
            _set_cell_bg(cell, bg)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run  = para.add_run(str(value))
            run.font.color.rgb = AUDI_DARK
            run.font.name      = "Calibri"
            run.font.size      = Pt(10)

    doc.add_paragraph()  # spacing after table


# ── Images ────────────────────────────────────────────────────────────────────

def add_image_safe(
    doc: Document,
    image_path: str,
    width: float = 5.5,
    caption: str | None = None,
    align=WD_ALIGN_PARAGRAPH.CENTER,
) -> bool:
    """
    Safely add an image paragraph. Returns True if successful, False if skipped.

    Always use this instead of doc.add_picture() directly, to handle corrupt
    or missing files gracefully.
    """
    abs_path = Path(image_path)
    if not abs_path.exists():
        print(f"[docx_helper] Skipping missing image: {image_path}")
        return False
    if abs_path.stat().st_size == 0:
        print(f"[docx_helper] Skipping empty image: {image_path}")
        return False

    valid_exts = {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp"}
    if abs_path.suffix.lower() not in valid_exts:
        print(f"[docx_helper] Skipping unsupported format: {image_path}")
        return False

    try:
        p = doc.add_paragraph()
        p.alignment = align
        run = p.add_run()
        run.add_picture(str(abs_path), width=Inches(width))
    except Exception as exc:
        print(f"[docx_helper] Failed to add image '{image_path}': {exc}")
        return False

    if caption:
        cap_para = doc.add_paragraph(caption)
        cap_para.alignment = align
        cap_run  = cap_para.runs[0] if cap_para.runs else cap_para.add_run(caption)
        cap_run.font.size  = Pt(9)
        cap_run.font.color.rgb = AUDI_SILVER
        cap_run.italic     = True

    return True


# ── Text helpers ──────────────────────────────────────────────────────────────

def add_body_text(
    doc: Document,
    text: str,
    font_size: int = 11,
    color: RGBColor = AUDI_DARK,
    bold: bool = False,
    space_after: float = 6,
) -> None:
    """Add a regular body paragraph."""
    p   = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold           = bold
    run.font.size      = Pt(font_size)
    run.font.color.rgb = color
    run.font.name      = "Calibri"


def add_bullet(doc: Document, text: str, font_size: int = 11) -> None:
    """Add a bullet-list paragraph (List Bullet style)."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size      = Pt(font_size)
    run.font.color.rgb = AUDI_DARK
    run.font.name      = "Calibri"
